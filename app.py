import os
import json
import sqlite3
from flask import Flask, request, jsonify, render_template, session, redirect, url_for
from werkzeug.security import generate_password_hash, check_password_hash
import PyPDF2
from docx import Document
import io
from flask import send_file
from fpdf import FPDF

import database
import ai_service

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "ai_insight_extractor_static_secret_key_123")
app.config.update(
    SESSION_COOKIE_SAMESITE='None',
    SESSION_COOKIE_SECURE=True
)

# Ensure db initialized
database.init_db()

# Safe Migration: Add share_id column if it doesn't exist for older documents
try:
    conn = database.get_db_connection()
    conn.execute("ALTER TABLE documents ADD COLUMN share_id TEXT")
    conn.commit()
    conn.close()
except sqlite3.OperationalError:
    pass # Column already exists

# ---- Authentication Routes ---- #
@app.route("/api/auth/register", methods=["POST"])
def register():
    data = request.json
    email = data.get("email")
    password = data.get("password")
    if not email or not password:
        return jsonify({"error": "Missing email or password"}), 400
    
    conn = database.get_db_connection()
    try:
        conn.execute("INSERT INTO users (email, password) VALUES (?, ?)", 
                 (email, generate_password_hash(password)))
        conn.commit()
    except sqlite3.IntegrityError:
        return jsonify({"error": "Email already exists"}), 400
    finally:
        conn.close()
    return jsonify({"success": True, "message": "Registered successfully"})

@app.route("/api/auth/login", methods=["POST"])
def login():
    data = request.json
    email = data.get("email")
    password = data.get("password")
    
    conn = database.get_db_connection()
    user = conn.execute("SELECT * FROM users WHERE email = ?", (email,)).fetchone()
    conn.close()
    
    if user and check_password_hash(user["password"], password):
        session["user_id"] = user["id"]
        return jsonify({"success": True})
    return jsonify({"error": "Invalid credentials"}), 401

@app.route("/api/auth/logout", methods=["POST"])
def logout():
    session.pop("user_id", None)
    return jsonify({"success": True})

# ---- Core Routes ---- #
@app.route("/")
def index():
    if "user_id" in session:
        return redirect(url_for("dashboard"))
    return render_template("index.html")

@app.route("/dashboard")
def dashboard():
    if "user_id" not in session:
        return redirect(url_for("index"))
    return render_template("dashboard.html")

@app.route("/analysis/<int:doc_id>")
def analysis(doc_id):
    if "user_id" not in session:
        return redirect(url_for("index"))
    return render_template("analysis.html", doc_id=doc_id)

@app.route("/about")
def about():
    return render_template("about.html")

@app.route("/contact")
def contact():
    return render_template("contact.html")

# ---- File Upload ---- #
def extract_text(file):
    filename = file.filename.lower()
    text = ""
    try:
        if filename.endswith(".pdf"):
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        elif filename.endswith(".docx"):
            doc = Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        elif filename.endswith(".txt"):
            text = file.read().decode("utf-8")
    except Exception as e:
        print(f"Extraction error: {e}")
    return text.strip()

@app.route("/api/upload", methods=["POST"])
def upload_file():
    if "user_id" not in session:
        return jsonify({"error": "Unauthorized"}), 401
        
    if "file" not in request.files:
        return jsonify({"error": "No file part"}), 400
        
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No selected file"}), 400
        
    text = extract_text(file)
    if not text:
        return jsonify({"error": "Could not extract text from file"}), 400
        
    conn = database.get_db_connection()
    cursor = conn.cursor()
    cursor.execute("INSERT INTO documents (user_id, filename, content) VALUES (?, ?, ?)", 
                   (session["user_id"], file.filename, text))
    conn.commit()
    doc_id = cursor.lastrowid
    conn.close()
    
    return jsonify({"success": True, "doc_id": doc_id})

# ---- AI Endpoints ---- #
@app.route("/api/documents", methods=["GET"])
def get_documents():
    if "user_id" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    conn = database.get_db_connection()
    docs = conn.execute("SELECT id, filename, upload_date FROM documents WHERE user_id = ? ORDER BY id DESC", 
                        (session["user_id"],)).fetchall()
    conn.close()
    return jsonify({"documents": [dict(d) for d in docs]})

@app.route("/api/documents/<int:doc_id>", methods=["DELETE"])
def delete_document(doc_id):
    if "user_id" not in session:
        return jsonify({"error": "Unauthorized"}), 401
        
    conn = database.get_db_connection()
    doc = conn.execute("SELECT id FROM documents WHERE id = ? AND user_id = ?", 
                       (doc_id, session["user_id"])).fetchone()
    if not doc:
        conn.close()
        return jsonify({"error": "Document not found"}), 404
        
    cursor = conn.cursor()
    cursor.execute("DELETE FROM summaries WHERE doc_id = ?", (doc_id,))
    cursor.execute("DELETE FROM chats WHERE doc_id = ?", (doc_id,))
    cursor.execute("DELETE FROM mcqs WHERE doc_id = ?", (doc_id,))
    cursor.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
    conn.commit()
    conn.close()
    
    return jsonify({"success": True})

@app.route("/api/summarize/<int:doc_id>", methods=["POST"])
def summarize(doc_id):
    if "user_id" not in session:
        return jsonify({"error": "Unauthorized"}), 401
        
    data = request.json
    size = data.get("size", "medium")
    
    conn = database.get_db_connection()
    doc = conn.execute("SELECT content FROM documents WHERE id = ? AND user_id = ?", 
                       (doc_id, session["user_id"])).fetchone()
    
    if not doc:
        conn.close()
        return jsonify({"error": "Document not found"}), 404
        
    # Check if summary exists for this size
    existing = conn.execute("SELECT * FROM summaries WHERE doc_id = ? AND size = ?", 
                            (doc_id, size)).fetchone()
    if existing:
        conn.close()
        return jsonify(dict(existing))
        
    # Call Gemini
    result = ai_service.get_summary_and_insights(doc["content"], size)
    
    # Convert lists to strings for SQLite
    def format_field(val):
        if isinstance(val, list):
            return "\\n".join(f"- {v}" for v in val)
        return str(val) if val is not None else ""

    try:
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO summaries (doc_id, size, summary_text, context, methodology, key_features, advantages, disadvantages)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ''', (doc_id, size, format_field(result.get("summary")), format_field(result.get("context")), 
              format_field(result.get("methodology")), format_field(result.get("key_features")),
              format_field(result.get("advantages")), format_field(result.get("disadvantages"))))
        conn.commit()
        summary_id = cursor.lastrowid
        res = dict(conn.execute("SELECT * FROM summaries WHERE id = ?", (summary_id,)).fetchone())
        conn.close()
        return jsonify(res)
    except Exception as e:
        conn.close()
        print(f"DB Error: {e}")
        return jsonify({"error": "Failed to save summary"}), 500

@app.route("/api/download/<format>/<int:doc_id>")
def download_summary(format, doc_id):
    if "user_id" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    
    conn = database.get_db_connection()
    doc = conn.execute("SELECT filename FROM documents WHERE id = ? AND user_id = ?", 
                       (doc_id, session["user_id"])).fetchone()
    if not doc:
        conn.close()
        return jsonify({"error": "Unauthorized"}), 404
        
    size = request.args.get("size", "medium")
    summary = conn.execute("SELECT * FROM summaries WHERE doc_id = ? AND size = ?", 
                           (doc_id, size)).fetchone()
    conn.close()
    
    if not summary:
        return jsonify({"error": "Summary not found"}), 404
        
    base_name = os.path.splitext(doc["filename"])[0]
    
    # helper text generator
    def build_text():
        lines = []
        lines.append(f"Analysis for: {doc['filename']}")
        lines.append("="*40)
        lines.append("\\n--- Main Summary ---\\n" + summary["summary_text"])
        lines.append("\\n--- Context ---\\n" + str(summary["context"]))
        lines.append("\\n--- Methodology ---\\n" + str(summary["methodology"]))
        lines.append("\\n--- Key Features ---\\n" + str(summary["key_features"]))
        lines.append("\\n--- Advantages ---\\n" + str(summary["advantages"]))
        lines.append("\\n--- Disadvantages ---\\n" + str(summary["disadvantages"]))
        
        text = "\\n".join(lines)
        # Sanitize smart quotes and dashes for latin-1 FPDF limitations
        replacements = {
            '—': '-', '–': '-',
            '“': '"', '”': '"',
            '‘': "'", '’': "'"
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        return text.encode('latin-1', 'replace').decode('latin-1')

    if format == 'txt':
        mem = io.BytesIO()
        mem.write(build_text().encode('utf-8'))
        mem.seek(0)
        return send_file(mem, as_attachment=True, download_name=f"{base_name}_summary.txt", mimetype='text/plain')
        
    elif format == 'docx':
        from docx import Document as DocxDoc
        d = DocxDoc()
        d.add_heading(f"Analysis for: {doc['filename']}", 0)
        d.add_heading("Main Summary", level=1)
        d.add_paragraph(summary["summary_text"])
        d.add_heading("Context", level=1)
        d.add_paragraph(str(summary["context"]))
        d.add_heading("Methodology", level=1)
        d.add_paragraph(str(summary["methodology"]))
        d.add_heading("Key Features", level=1)
        d.add_paragraph(str(summary["key_features"]))
        d.add_heading("Advantages", level=1)
        d.add_paragraph(str(summary["advantages"]))
        d.add_heading("Disadvantages", level=1)
        d.add_paragraph(str(summary["disadvantages"]))
        
        mem = io.BytesIO()
        d.save(mem)
        mem.seek(0)
        return send_file(mem, as_attachment=True, download_name=f"{base_name}_summary.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
    elif format == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, txt=build_text())
        
        out = pdf.output(dest='S')
        if isinstance(out, (bytes, bytearray)):
            mem = io.BytesIO(out)
        else:
            mem = io.BytesIO(out.encode('latin-1', 'replace'))
            
        return send_file(mem, as_attachment=True, download_name=f"{base_name}_summary.pdf", mimetype='application/pdf')

    return jsonify({"error": "Invalid format"}), 400

@app.route("/api/chat/<int:doc_id>", methods=["POST", "GET"])
def chat(doc_id):
    if "user_id" not in session:
        return jsonify({"error": "Unauthorized"}), 401
        
    conn = database.get_db_connection()
    doc = conn.execute("SELECT content FROM documents WHERE id = ? AND user_id = ?", 
                       (doc_id, session["user_id"])).fetchone()
    if not doc:
        conn.close()
        return jsonify({"error": "Doc not found"}), 404
        
    if request.method == "GET":
        chats = conn.execute("SELECT role, message FROM chats WHERE doc_id = ? ORDER BY id ASC", 
                             (doc_id,)).fetchall()
        conn.close()
        return jsonify({"chats": [dict(c) for c in chats]})
        
    # POST - New message
    data = request.json
    user_message = data.get("message", "")
    if not user_message:
         return jsonify({"error": "Message is empty"}), 400
         
    # Save user message
    cursor = conn.cursor()
    cursor.execute("INSERT INTO chats (doc_id, role, message) VALUES (?, ?, ?)", 
                   (doc_id, "user", user_message))
                   
    # Get history
    history = conn.execute("SELECT role, message FROM chats WHERE doc_id = ? ORDER BY id ASC", (doc_id,)).fetchall()
    
    # Get AI response
    ai_response = ai_service.answer_question(doc["content"], [dict(h) for h in history[:-1]], user_message)
    
    # Save AI message
    cursor.execute("INSERT INTO chats (doc_id, role, message) VALUES (?, ?, ?)", 
                   (doc_id, "ai", ai_response))
    conn.commit()
    conn.close()
    
    return jsonify({"response": ai_response})

@app.route("/api/mcq/<int:doc_id>", methods=["POST", "GET"])
def mcq(doc_id):
    if "user_id" not in session:
        return jsonify({"error": "Unauthorized"}), 401
        
    conn = database.get_db_connection()
    
    if request.method == "GET":
        existing = conn.execute("SELECT questions FROM mcqs WHERE doc_id = ?", (doc_id,)).fetchone()
        conn.close()
        if existing:
            return jsonify({"mcqs": json.loads(existing["questions"])})
        return jsonify({"mcqs": []})
        
    # POST - Generate
    doc = conn.execute("SELECT content FROM documents WHERE id = ? AND user_id = ?", 
                       (doc_id, session["user_id"])).fetchone()
    if not doc:
        conn.close()
        return jsonify({"error": "Doc not found"}), 404
        
    questions = ai_service.generate_mcqs(doc["content"])
    
    if questions:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM mcqs WHERE doc_id = ?", (doc_id,))
        cursor.execute("INSERT INTO mcqs (doc_id, questions) VALUES (?, ?)", 
                       (doc_id, json.dumps(questions)))
        conn.commit()
    
    conn.close()
    return jsonify({"mcqs": questions})

@app.route("/api/mcq/score/<int:doc_id>", methods=["POST", "GET"])
def mcq_score(doc_id):
    if "user_id" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    conn = database.get_db_connection()
    
    if request.method == "GET":
        scores = conn.execute("SELECT * FROM test_scores WHERE doc_id = ? ORDER BY id DESC", (doc_id,)).fetchall()
        conn.close()
        return jsonify({"scores": [dict(s) for s in scores]})
        
    data = request.json
    score_val = data.get("score")
    total_val = data.get("total")
    
    if score_val is not None and total_val is not None:
        cursor = conn.cursor()
        cursor.execute("INSERT INTO test_scores (doc_id, score, total) VALUES (?, ?, ?)", 
                       (doc_id, score_val, total_val))
        conn.commit()
        conn.close()
        return jsonify({"success": True})
        
    conn.close()
    return jsonify({"error": "Invalid data"}), 400

@app.route("/api/tret-ai/chat", methods=["POST"])
def tret_ai_chat():
    """ Stateless general chat endpoint that trusts the history passed by frontend. """
    data = request.json
    history = data.get("history", []) # Array of {role: 'user'|'ai', message: string}
    new_message = data.get("message", "")
    
    if not new_message:
        return jsonify({"error": "Message is empty"}), 400
        
    ai_response = ai_service.tret_ai_general_chat(history, new_message)
    return jsonify({"response": ai_response})

# ---- Next-Gen Feature Routes ---- #

import uuid

@app.route("/api/flashcards/<int:doc_id>", methods=["POST", "GET"])
def flashcards(doc_id):
    if "user_id" not in session: return jsonify({"error": "Unauthorized"}), 401
    conn = database.get_db_connection()
    if request.method == "GET":
        existing = conn.execute("SELECT flashcards FROM flashcards WHERE doc_id = ?", (doc_id,)).fetchone()
        conn.close()
        if existing: return jsonify({"flashcards": json.loads(existing["flashcards"])})
        return jsonify({"flashcards": []})
        
    doc = conn.execute("SELECT content FROM documents WHERE id = ? AND user_id = ?", (doc_id, session["user_id"])).fetchone()
    if not doc:
        conn.close()
        return jsonify({"error": "Doc not found"}), 404
        
    cards = ai_service.generate_flashcards(doc["content"])
    if cards:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM flashcards WHERE doc_id = ?", (doc_id,))
        cursor.execute("INSERT INTO flashcards (doc_id, flashcards) VALUES (?, ?)", (doc_id, json.dumps(cards)))
        conn.commit()
    conn.close()
    return jsonify({"flashcards": cards})

@app.route("/api/download/csv/flashcards/<int:doc_id>")
def download_flashcards_csv(doc_id):
    if "user_id" not in session: return jsonify({"error": "Unauthorized"}), 401
    conn = database.get_db_connection()
    doc = conn.execute("SELECT filename FROM documents WHERE id = ? AND user_id = ?", (doc_id, session["user_id"])).fetchone()
    existing = conn.execute("SELECT flashcards FROM flashcards WHERE doc_id = ?", (doc_id,)).fetchone()
    conn.close()
    
    if not existing or not doc: return jsonify({"error": "Not found"}), 404
    cards = json.loads(existing["flashcards"])
    
    import csv
    mem = io.StringIO()
    writer = csv.writer(mem)
    writer.writerow(["Front", "Back"])
    for c in cards:
        writer.writerow([c.get("front", ""), c.get("back", "")])
    
    mem.seek(0)
    bytes_mem = io.BytesIO(mem.getvalue().encode('utf-8'))
    return send_file(bytes_mem, as_attachment=True, download_name=f"{doc['filename']}_flashcards.csv", mimetype='text/csv')

@app.route("/api/share/<int:doc_id>", methods=["POST"])
def enable_share(doc_id):
    if "user_id" not in session: return jsonify({"error": "Unauthorized"}), 401
    conn = database.get_db_connection()
    doc = conn.execute("SELECT share_id FROM documents WHERE id = ? AND user_id = ?", (doc_id, session["user_id"])).fetchone()
    if not doc:
        conn.close()
        return jsonify({"error": "Doc not found"}), 404
        
    if doc["share_id"]:
        conn.close()
        return jsonify({"share_id": doc["share_id"]})
        
    share_uuid = str(uuid.uuid4())
    conn.execute("UPDATE documents SET share_id = ? WHERE id = ?", (share_uuid, doc_id))
    conn.commit()
    conn.close()
    return jsonify({"share_id": share_uuid})

@app.route("/share/<share_id>")
def view_shared(share_id):
    if request.remote_addr not in ['127.0.0.1', '::1', 'localhost']:
        return "Access denied: Share link only works on the system it was created from.", 403
    conn = database.get_db_connection()
    doc = conn.execute("SELECT id FROM documents WHERE share_id = ?", (share_id,)).fetchone()
    conn.close()
    if not doc:
        return "Shared document not found or link expired.", 404
    return render_template("shared_analysis.html", share_id=share_id)

@app.route("/api/public/data/<share_id>")
def public_data(share_id):
    if request.remote_addr not in ['127.0.0.1', '::1', 'localhost']:
        return jsonify({"error": "Access denied"}), 403
    conn = database.get_db_connection()
    doc = conn.execute("SELECT id, filename FROM documents WHERE share_id = ?", (share_id,)).fetchone()
    if not doc:
        conn.close()
        return jsonify({"error": "Not found"}), 404
        
    doc_id = doc["id"]
    summary = conn.execute("SELECT * FROM summaries WHERE doc_id = ? AND size = 'medium'", (doc_id,)).fetchone()
    mcq = conn.execute("SELECT questions FROM mcqs WHERE doc_id = ?", (doc_id,)).fetchone()
    flashcards = conn.execute("SELECT flashcards FROM flashcards WHERE doc_id = ?", (doc_id,)).fetchone()
    conn.close()
    
    return jsonify({
        "filename": doc["filename"],
        "summary": dict(summary) if summary else None,
        "mcqs": json.loads(mcq["questions"]) if mcq else [],
        "flashcards": json.loads(flashcards["flashcards"]) if flashcards else []
    })

@app.route("/api/compare", methods=["POST"])
def compare_docs():
    if "user_id" not in session: return jsonify({"error": "Unauthorized"}), 401
    doc_ids = request.json.get("doc_ids", [])
    if len(doc_ids) < 2: return jsonify({"error": "Need at least 2 documents to compare"}), 400
    
    conn = database.get_db_connection()
    placeholders = ",".join("?" for _ in doc_ids)
    query = f"SELECT content FROM documents WHERE user_id = ? AND id IN ({placeholders})"
    params = [session["user_id"]] + doc_ids
    docs = conn.execute(query, params).fetchall()
    conn.close()
    
    if len(docs) != len(doc_ids):
        return jsonify({"error": "One or more documents could not be accessed"}), 404
        
    texts = [d["content"] for d in docs]
    result = ai_service.compare_documents(texts)
    return jsonify(result)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=7860)
